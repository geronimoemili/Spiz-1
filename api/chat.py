"""
api/chat.py — SPIZ Intelligence Core

FUNZIONALITA':
- Conta/autori/testate/AVE → Supabase diretto, 0 token OpenAI
- Analisi/report testuale → OpenAI con ricerca keyword
- Report Word professionale → fetch articoli + GPT-4o + python-docx
- Memoria conversazionale per sessione
- Multi-cliente
"""

import os
import re
import json
import uuid
import openai
import tiktoken
from datetime import date, timedelta
from pathlib import Path
from services.database import supabase

# ─── CONFIG ────────────────────────────────────────────────────────────────────
MODEL_SMART         = "gpt-4o"
MODEL_FAST          = "gpt-4o-mini"
MAX_CONTEXT_TOKENS  = 12_000
MAX_RESPONSE_TOKENS = 1_000
BASE_URL            = os.getenv("APP_BASE_URL", "https://tua-app.replit.app")

# Cartella dove vengono salvati i report Word generati
REPORTS_DIR = Path("static/reports")
REPORTS_DIR.mkdir(parents=True, exist_ok=True)

enc = tiktoken.encoding_for_model("gpt-4o")

def _tokens(text: str) -> int:
    return len(enc.encode(text))

# ─── MEMORIA CONVERSAZIONE ──────────────────────────────────────────────────────
_sessions: dict[str, list[dict]] = {}

def get_history(session_id: str) -> list[dict]:
    return _sessions.setdefault(session_id, [])

def save_turn(session_id: str, question: str, answer: str):
    h = get_history(session_id)
    h.append({"role": "user",      "content": question})
    h.append({"role": "assistant", "content": answer})

def trim_history(history: list[dict], budget: int = 3_000) -> list[dict]:
    out, used = [], 0
    for msg in reversed(history):
        t = _tokens(msg["content"])
        if used + t > budget:
            break
        out.insert(0, msg)
        used += t
    return out

def reset_session(session_id: str):
    _sessions[session_id] = []

# ─── CLIENTI ───────────────────────────────────────────────────────────────────
def load_client(client_id: str) -> dict | None:
    try:
        res = supabase.table("clients").select("*").eq("id", client_id).execute()
        if res.data:
            return res.data[0]
        res = supabase.table("clients").select("*").ilike("name", client_id).execute()
        return res.data[0] if res.data else None
    except Exception:
        return None

def list_clients() -> list[dict]:
    try:
        res = supabase.table("clients").select("id, name, keywords, semantic_topic").execute()
        return res.data or []
    except Exception:
        return []

def parse_keywords(raw: str) -> list[str]:
    if not raw:
        return []
    sep = "," if "," in raw else "\n"
    return [k.strip().lower() for k in raw.split(sep) if k.strip()]

# ─── RANGE TEMPORALE ───────────────────────────────────────────────────────────
def date_range_from_question(question: str) -> tuple[str, str]:
    q     = question.lower()
    today = date.today()

    if "ieri" in q:
        d = today - timedelta(days=1)
        return str(d), str(d)
    if "settimana scorsa" in q:
        end   = today - timedelta(days=today.weekday() + 1)
        start = end - timedelta(days=6)
        return str(start), str(end)
    if "questa settimana" in q or "settimana" in q:
        return str(today - timedelta(days=today.weekday())), str(today)
    if "questo mese" in q or "mese" in q:
        return str(today.replace(day=1)), str(today)
    m = re.search(r'ultim[oi]\s*(\d+)\s*giorn', q)
    if m:
        return str(today - timedelta(days=int(m.group(1)))), str(today)
    if "ultimi 7" in q or "7 giorni" in q:
        return str(today - timedelta(days=7)), str(today)
    if "ultimi 30" in q or "30 giorni" in q:
        return str(today - timedelta(days=30)), str(today)
    return str(today), str(today)

# ─── KEYWORD TOPIC ─────────────────────────────────────────────────────────────
_STOPWORDS = {
    "di","il","la","lo","le","i","un","una","è","e","per","che","con","da","in",
    "su","non","mi","si","ho","ha","hai","sono","fare","del","della","degli","delle",
    "dei","gli","questo","questa","questi","queste","qual","quale","quali","dammi",
    "voglio","mostra","analizza","cerca","quanti","quante","chi","cosa","come",
    "quando","dove","perché","tutti","tutte","tutto","anche","però","ancora","già",
    "sempre","mai","molto","poco","tanti","tante","alcuni","alcune","ogni","oppure",
    "oggi","ieri","settimana","mese","articoli","articolo","news","elenco","lista",
    "hanno","scritto","firma","firmati","senza","quelli","quali","prepara","fammi",
    "report","rassegna","quello","uscito","uscita","completo","completa","genera",
    "word","documento","scarica","crea","costruisci","dammi","voglio","deve",
    "giornalisti","giornalista","autori","autore","firme","testate","testata",
    "fonti","fonte","giornali","giornale","database","archivio","conteggio","conta",
}

def topic_keywords_from_question(question: str) -> list[str]:
    clean = re.sub(r"[^\w\s]", " ", question.lower())
    words = []
    SHORT_ALLOWED = {
        "mps","eni","enel","tim","ubi","bpm","bce","fed","irs","ipo","pil","pnrr",
        "a2a","gse","gse","erg","res","gas","oil","llp","spa","srl","ceo","cda",
    }
    for w in clean.split():
        if w in SHORT_ALLOWED:
            words.append(w)
        elif len(w) > 3 and w not in _STOPWORDS:
            words.append(w)
    return words[:6]

# ─── CLASSIFICAZIONE INTENT ─────────────────────────────────────────────────────
INTENTS = {
    "report_word": [
        "genera report", "crea report", "report word", "documento word",
        "scarica report", "report professionale", "report completo word",
        "genera documento", "esporta report", "voglio il report",
        "dammi il report", "report in word", "genera il report",
        "crea il documento", "report strutturato",
    ],
    "conta": [
        "quanti", "conta ", "conteggio", "numero di", "quante volte",
        "quante testate", "quanti giornalisti", "quanti articoli", "quante notizie",
        "firmati", "non firmati", "senza firma", "con firma", "anonimi",
    ],
    "ave": [
        "ave", "valore economico", "copertura economica", "quanto vale",
    ],
    "autore": [
        "chi ha scritto", "giornalista", "giornalisti", "autore", "autori",
        "firma", "firme", "chi parla", "elenco giornalisti", "elenco autori",
        "lista giornalisti", "lista autori", "chi scrive", "chi ha pubblicato",
    ],
    "fonte": [
        "quali testate", "quale giornale", "quali fonti", "elenco testate",
        "lista testate", "da quali giornali",
    ],
    "leggi": [
        "leggi", "mostrami il testo", "testo integrale", "voglio leggere",
        "testo completo", "articolo intero",
    ],
    "analisi": [
        "analizza", "analisi", "confronta", "trend", "andamento",
        "sintesi", "riassumi", "paragona", "cosa emerge", "cosa dicono",
    ],
    "report": [
        "report", "rassegna", "panoramica", "dammi un quadro", "sommario", "briefing",
    ],
}

def classify_intent(question: str) -> str:
    q = question.lower()
    for intent, triggers in INTENTS.items():
        if any(t in q for t in triggers):
            return intent
    return "generico"

# ─── FILTRI DB ─────────────────────────────────────────────────────────────────
def extract_db_filters(question: str) -> dict:
    q, f = question.lower(), {}
    if "negativ" in q:   f["tone"] = "Negative"
    elif "positiv" in q: f["tone"] = "Positive"
    elif "neutro" in q:  f["tone"] = "Neutral"
    if "rischio alto" in q or "alto rischio" in q:
        f["reputational_risk"] = "Alto"
    elif "rischio medio" in q:
        f["reputational_risk"] = "Medio"
    return f

# ─── CAMPI DB ─────────────────────────────────────────────────────────────────
_SEARCH_FIELDS = ["titolo", "testo_completo", "macrosettori", "dominant_topic", "occhiello"]

FIELDS_LIGHT = (
    "id, testata, data, titolo, giornalista, autore, "
    "tone, reputational_risk, dominant_topic, macrosettori, ave"
)

FIELDS_REPORT = (
    "id, testata, data, titolo, giornalista, autore, occhiello, "
    "tone, reputational_risk, political_risk, dominant_topic, "
    "macrosettori, tipologia_articolo, ave, tipo_fonte"
)

FIELDS_FULL = (
    "id, testata, data, titolo, giornalista, autore, occhiello, sottotitolo, "
    "testo_completo, tone, dominant_topic, reputational_risk, political_risk, "
    "macrosettori, tipologia_articolo, ave, tipo_fonte"
)

# ─── QUERY GENERICA ─────────────────────────────────────────────────────────────
def _run_query(fields: str, date_from: str, date_to: str, keywords: list[str],
               filters: dict, limit: int) -> list[dict]:
    seen, results = set(), []

    def _add(rows):
        for r in (rows or []):
            if r["id"] not in seen:
                seen.add(r["id"])
                results.append(r)

    if keywords:
        for kw in keywords:
            for field in _SEARCH_FIELDS:
                try:
                    q = (supabase.table("articles").select(fields)
                         .gte("data", date_from).lte("data", date_to)
                         .ilike(field, f"%{kw}%"))
                    for col, val in filters.items():
                        q = q.eq(col, val)
                    _add(q.limit(limit).execute().data)
                except Exception:
                    pass
    else:
        try:
            q = (supabase.table("articles").select(fields)
                 .gte("data", date_from).lte("data", date_to)
                 .order("data", desc=True))
            for col, val in filters.items():
                q = q.eq(col, val)
            _add(q.limit(limit).execute().data)
        except Exception:
            pass

    return results

# ─── RESOLVER DIRETTI (0 token OpenAI) ─────────────────────────────────────────
def _autore_str(r: dict) -> str:
    return (r.get("giornalista") or r.get("autore") or "").strip()

def _is_firmato(r: dict) -> bool:
    return _autore_str(r) not in ("", "N.D.", "N/D", "Redazione", "Autore non indicato")

def resolve_conta(rows: list[dict], question: str, date_from: str, date_to: str) -> str:
    if not rows:
        return "Nessun articolo trovato per i criteri specificati."
    periodo = f"del {date_from}" if date_from == date_to else f"dal {date_from} al {date_to}"
    q = question.lower()
    if any(t in q for t in ("firma", "firmati", "non firmati", "senza firma", "con firma", "anonimo")):
        firmati  = [r for r in rows if _is_firmato(r)]
        non_firm = [r for r in rows if not _is_firmato(r)]
        pct = len(firmati) / len(rows) * 100
        return (f"**Articoli {periodo}: {len(rows)} totali**\n\n"
                f"Con firma: {len(firmati)}\nSenza firma / Redazione: {len(non_firm)}\n"
                f"Firmati: {pct:.1f}%")
    by_testata: dict[str, int] = {}
    for r in rows:
        t = r.get("testata", "N/D")
        by_testata[t] = by_testata.get(t, 0) + 1
    table = "| Testata | Articoli |\n|---------|----------|\n"
    for testata, n in sorted(by_testata.items(), key=lambda x: -x[1]):
        table += f"| {testata} | {n} |\n"
    return (f"**Articoli trovati {periodo}: {len(rows)}**\n\n{table}\n"
            f"**Totale: {len(rows)} articoli su {len(by_testata)} testate**")

def resolve_autore(rows: list[dict], date_from: str, date_to: str) -> str:
    if not rows:
        return "Nessun articolo trovato."
    periodo = f"del {date_from}" if date_from == date_to else f"dal {date_from} al {date_to}"
    by_autore: dict[str, list[str]] = {}
    for r in rows:
        autore = _autore_str(r) or "Senza firma"
        entry  = f"*{r.get('testata','N/D')}* — {r.get('titolo','N/D')}"
        by_autore.setdefault(autore, []).append(entry)
    lines = [f"**Giornalisti {periodo} ({len(rows)} articoli | {len(by_autore)} firme):**\n"]
    for autore, titoli in sorted(by_autore.items(), key=lambda x: -len(x[1])):
        lines.append(f"\n**{autore}** ({len(titoli)} art.)")
        for t in titoli[:5]:
            lines.append(f"  • {t}")
        if len(titoli) > 5:
            lines.append(f"  *...e altri {len(titoli)-5}*")
    return "\n".join(lines)

def resolve_fonte(rows: list[dict], date_from: str, date_to: str) -> str:
    if not rows:
        return "Nessun articolo trovato."
    periodo = f"del {date_from}" if date_from == date_to else f"dal {date_from} al {date_to}"
    by_testata: dict[str, list[dict]] = {}
    for r in rows:
        by_testata.setdefault(r.get("testata", "N/D"), []).append(r)
    lines = [f"**Testate {periodo} ({len(rows)} articoli totali):**\n"]
    for testata, arts in sorted(by_testata.items(), key=lambda x: -len(x[1])):
        autori = {_autore_str(a) for a in arts if _is_firmato(a)}
        lines.append(f"\n**{testata}** — {len(arts)} articoli")
        if autori:
            lines.append(f"  Firme: {', '.join(sorted(autori))}")
    return "\n".join(lines)

def resolve_ave(rows: list[dict], date_from: str, date_to: str) -> str:
    if not rows:
        return "Nessun articolo trovato."
    periodo = f"del {date_from}" if date_from == date_to else f"dal {date_from} al {date_to}"
    table = "| Testata | Titolo | AVE (€) |\n|---------|--------|---------|\n"
    totale, senza = 0.0, 0
    for r in sorted(rows, key=lambda x: float(x.get("ave") or 0), reverse=True)[:50]:
        testata = r.get("testata", "N/D")
        titolo  = (r.get("titolo") or "N/D")[:55]
        try:
            ave_val = float(r.get("ave"))
            totale += ave_val
            table += f"| {testata} | {titolo} | {ave_val:,.0f} |\n"
        except (TypeError, ValueError):
            senza += 1
            table += f"| {testata} | {titolo} | N/D |\n"
    footer = f"\n**AVE TOTALE {periodo}: €{totale:,.0f}**"
    if senza:
        footer += f"\n*{senza} articoli senza AVE esclusi.*"
    return table + footer

# ─── FETCH PER OPENAI (con testo) ──────────────────────────────────────────────
def fetch_full_articles(keywords: list[str], date_from: str, date_to: str,
                        extra_filters: dict | None = None, limit: int = 8) -> list[dict]:
    results = _run_query(FIELDS_FULL, date_from, date_to, keywords, extra_filters or {}, limit)
    return results[:limit]

def build_context(articles: list[dict], budget: int) -> tuple[str, int]:
    context, used, n = "", 0, 0
    for a in articles:
        link = f"{BASE_URL}/articolo/{a.get('id','')}"
        chunk = (
            f"=== ARTICOLO [{n+1}] ===\n"
            f"TESTATA: {a.get('testata','N/D')} | DATA: {a.get('data','N/D')}\n"
            f"AUTORE: {a.get('giornalista') or a.get('autore','N/D')}\n"
            f"TITOLO: {a.get('titolo','N/D')}\n"
            f"TONE: {a.get('tone','N/D')} | RISCHIO: {a.get('reputational_risk','N/D')}\n"
            f"SETTORI: {a.get('macrosettori','N/D')}\n"
            f"AVE: {a.get('ave','N/D')}\n"
            f"LINK: {link}\n"
            f"TESTO:\n{a.get('testo_completo','')}\n\n"
        )
        t = _tokens(chunk)
        if used + t > budget:
            break
        context += chunk
        used += t
        n += 1
    return f"[{n}/{len(articles)} articoli | {used:,} token]\n\n" + context, n

# ═══════════════════════════════════════════════════════════════════════════════
#  GENERAZIONE REPORT WORD
# ═══════════════════════════════════════════════════════════════════════════════

def _build_report_context(articles: list[dict]) -> str:
    """
    Costruisce il contesto per GPT-4o con metadati + occhiello (no testo integrale).
    Piu efficiente in token rispetto al testo completo.
    """
    lines = []
    for i, a in enumerate(articles, 1):
        occhiello = (a.get("occhiello") or "")[:200]
        lines.append(
            f"[{i}] TESTATA: {a.get('testata','N/D')} | DATA: {a.get('data','N/D')} | "
            f"AVE: {a.get('ave','N/D')} | TONE: {a.get('tone','N/D')} | "
            f"RISCHIO_REP: {a.get('reputational_risk','N/D')} | "
            f"RISCHIO_POL: {a.get('political_risk','N/D')} | "
            f"SETTORI: {a.get('macrosettori','N/D')} | "
            f"TIPOLOGIA: {a.get('tipologia_articolo','N/D')} | "
            f"FONTE: {a.get('tipo_fonte','N/D')}\n"
            f"    AUTORE: {a.get('giornalista') or a.get('autore','N/D')}\n"
            f"    TITOLO: {a.get('titolo','N/D')}\n"
            f"    OCCHIELLO: {occhiello}\n"
        )
    return "\n".join(lines)


def _call_gpt_for_report(articles: list[dict], keywords: list[str],
                          client_data: dict | None, date_from: str, date_to: str) -> dict:
    """
    Chiama GPT-4o per generare le 10 sezioni del report in JSON.
    Usa solo metadati + occhiello per restare nei limiti di token.
    """
    periodo = f"del {date_from}" if date_from == date_to else f"dal {date_from} al {date_to}"
    cliente_str = client_data.get("name") if client_data else "Area generale"
    topic_str   = ", ".join(keywords) if keywords else "generale"
    ctx         = _build_report_context(articles)
    n_art       = len(articles)

    # Calcola AVE totale e sentiment dai dati
    ave_totale = sum(float(a.get("ave") or 0) for a in articles if a.get("ave"))
    pos = sum(1 for a in articles if str(a.get("tone","")).lower() in ("positive","positivo"))
    neg = sum(1 for a in articles if str(a.get("tone","")).lower() in ("negative","negativo"))
    neu = n_art - pos - neg

    # Testate uniche
    testate = sorted(set(a.get("testata","") for a in articles if a.get("testata")))
    autori_list = sorted(set(
        (a.get("giornalista") or a.get("autore","")).strip()
        for a in articles
        if (a.get("giornalista") or a.get("autore","")).strip()
        and (a.get("giornalista") or a.get("autore","")).strip() not in
            ("N/D","N.D.","Redazione","")
    ))

    system_prompt = """Sei un analista senior di comunicazione e media monitoring di un'agenzia professionale chiamata MAIM.
Devi generare un report di rassegna stampa in formato JSON strutturato.
Rispondi SOLO con un oggetto JSON valido, senza markdown, senza backtick, senza spiegazioni.
Il JSON deve avere esattamente queste chiavi:
{
  "sezione1_profilo": "testo analitico sulla narrativa mediatica emergente, ruoli attribuiti, percezione generale, temi associati",
  "sezione2_vertici": "testo su interviste, dichiarazioni, citazioni dei vertici presenti nella rassegna",
  "sezione3_temi_longevi": "testo sui temi ricorrenti raggruppati per macro-area, con indicazione persistenti/emergenti/in diminuzione",
  "sezione4_finanziario": "testo su notizie finanziarie e corporate rilevanti",
  "sezione5_governance": "testo su cambi di management, nomine, governance",
  "sezione6_territoriale": "testo su focus territoriale, aree geografiche citate, opposizioni locali, comitati",
  "sezione7_criticita": "testo sulle criticita reputazionali, trattazioni negative o potenzialmente dannose",
  "sezione8_sentiment": "testo sull'analisi del sentiment con percentuali, driver positivi e negativi",
  "sezione9_istituzionale": "testo sulla comunicazione istituzionale, rapporti con Governo/UE, ruolo strategico",
  "sezione10_sintesi": "testo con conclusioni operative: evoluzione percezione, territori da presidiare, priorita comunicative, rischi emergenti, opportunita narrative"
}
Ogni sezione deve essere un testo discorsivo professionale in italiano, approfondito, stile report per direzione comunicazione corporate.
Non usare emoji. Non usare markdown. Solo testo piano che sara poi formattato nel documento Word."""

    user_prompt = f"""Elabora il seguente report di rassegna stampa.

CLIENTE / AREA: {cliente_str}
TOPIC MONITORATO: {topic_str}
PERIODO: {periodo}
TOTALE ARTICOLI: {n_art}
TESTATE COINVOLTE ({len(testate)}): {', '.join(testate[:30])}
AVE TOTALE STIMATO: €{ave_totale:,.0f}
SENTIMENT: Positivi {pos} | Neutri {neu} | Negativi {neg}
PRINCIPALI AUTORI: {', '.join(autori_list[:20])}

ARTICOLI DELLA RASSEGNA:
{ctx}

Genera il report JSON con le 10 sezioni richieste. Sii approfondito e analitico."""

    response = openai.chat.completions.create(
        model=MODEL_SMART,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": user_prompt},
        ],
        temperature=0.3,
        max_tokens=4000,
    )

    raw = response.choices[0].message.content.strip()
    # Rimuovi eventuale markdown residuo
    raw = re.sub(r"^```json\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    return json.loads(raw)


def _build_word_document(sections: dict, articles: list[dict], keywords: list[str],
                          client_data: dict | None, date_from: str, date_to: str) -> Path:
    """
    Costruisce il documento Word professionale con python-docx.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import lxml.etree as etree

    # Colori
    C_DARK  = RGBColor(0x1A, 0x2E, 0x4A)
    C_MID   = RGBColor(0x2E, 0x6D, 0xA4)
    C_LIGHT = RGBColor(0xD6, 0xE8, 0xF7)
    C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    C_RED   = RGBColor(0xC0, 0x00, 0x00)
    C_GRAY  = RGBColor(0xF5, 0xF5, 0xF5)

    periodo = f"del {date_from}" if date_from == date_to else f"dal {date_from} al {date_to}"
    cliente_str = client_data.get("name") if client_data else "Area Energia"
    topic_str   = ", ".join(keywords).title() if keywords else "Rassegna Generale"

    # Statistiche
    n_art = len(articles)
    ave_totale = sum(float(a.get("ave") or 0) for a in articles if a.get("ave"))
    pos = sum(1 for a in articles if str(a.get("tone","")).lower() in ("positive","positivo"))
    neg = sum(1 for a in articles if str(a.get("tone","")).lower() in ("negative","negativo"))
    neu = n_art - pos - neg
    testate = sorted(set(a.get("testata","") for a in articles if a.get("testata")))

    doc = Document()

    # ── Impostazioni pagina ──
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    # ── Stili base ──
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Arial"
    style_normal.font.size = Pt(10)

    def set_heading_style(level: int, text: str):
        p = doc.add_heading(text, level=level)
        run = p.runs[0]
        run.font.name = "Arial"
        if level == 1:
            run.font.size = Pt(14)
            run.font.color.rgb = C_DARK
            run.bold = True
        elif level == 2:
            run.font.size = Pt(12)
            run.font.color.rgb = C_MID
            run.bold = True
        elif level == 3:
            run.font.size = Pt(11)
            run.font.color.rgb = C_DARK
            run.bold = True
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after  = Pt(6)
        return p

    def add_body(text: str):
        if not text:
            return
        # Spezza il testo in paragrafi sulle doppie righe
        for block in text.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(6)
            p.paragraph_format.space_before = Pt(2)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(block)
            run.font.name = "Arial"
            run.font.size = Pt(10)

    def add_bullet(text: str, bold_prefix: str = ""):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.space_before = Pt(1)
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.name = "Arial"
            rb.font.size = Pt(10)
            rb.bold = True
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(10)

    def shade_cell(cell, hex_color: str):
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), hex_color)
        cell._tc.get_or_add_tcPr().append(shading)

    def cell_text(cell, text: str, bold=False, font_size=9,
                  color: RGBColor = None, align=WD_ALIGN_PARAGRAPH.LEFT):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(str(text or ""))
        run.font.name = "Arial"
        run.font.size = Pt(font_size)
        run.bold = bold
        if color:
            run.font.color.rgb = color

    def add_separator():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        border = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "2E6DA4")
        border.append(bottom)
        p._p.get_or_add_pPr().append(border)

    # ════════════════════════════════════════════════════════════════
    # COPERTINA
    # ════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()

    p_agenzia = doc.add_paragraph()
    p_agenzia.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_agenzia.add_run("MAIM")
    r.font.name = "Arial"
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = C_DARK

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_sub.add_run("Intelligence & Media Monitoring")
    r.font.name  = "Arial"
    r.font.size  = Pt(13)
    r.font.italic = True
    r.font.color.rgb = C_MID

    doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run("REPORT DI ANALISI MEDIA")
    r.font.name  = "Arial"
    r.font.size  = Pt(18)
    r.font.bold  = True
    r.font.color.rgb = C_DARK

    p_client = doc.add_paragraph()
    p_client.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_client.add_run(f"{cliente_str} — {topic_str}")
    r.font.name  = "Arial"
    r.font.size  = Pt(14)
    r.font.bold  = True
    r.font.color.rgb = C_MID

    p_data = doc.add_paragraph()
    p_data.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_data.add_run(periodo.capitalize())
    r.font.name  = "Arial"
    r.font.size  = Pt(12)
    r.font.color.rgb = C_DARK

    doc.add_paragraph()

    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_info.add_run("Elaborato da: Direzione Analisi e Comunicazione — MAIM")
    r.font.name = "Arial"
    r.font.size = Pt(10)

    p_dest = doc.add_paragraph()
    p_dest.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_dest.add_run(f"Destinatari: Direzione Comunicazione — {cliente_str}")
    r.font.name = "Arial"
    r.font.size = Pt(10)

    p_riserv = doc.add_paragraph()
    p_riserv.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_riserv.add_run("Classificazione: Riservato")
    r.font.name  = "Arial"
    r.font.size  = Pt(10)
    r.font.bold  = True
    r.font.color.rgb = C_RED

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # SCHEDA QUANTITATIVA
    # ════════════════════════════════════════════════════════════════
    set_heading_style(1, "Scheda Quantitativa della Rassegna")

    tbl = doc.add_table(rows=2, cols=5)
    tbl.style = "Table Grid"
    hdrs = ["Articoli totali", "Testate coinvolte", "AVE totale (€)", "Positivi", "Negativi"]
    vals = [str(n_art), str(len(testate)), f"{ave_totale:,.0f}", str(pos), str(neg)]
    for i, (h, v) in enumerate(zip(hdrs, vals)):
        shade_cell(tbl.rows[0].cells[i], "1A2E4A")
        cell_text(tbl.rows[0].cells[i], h, bold=True, color=C_WHITE,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(tbl.rows[1].cells[i], "D6E8F7")
        cell_text(tbl.rows[1].cells[i], v, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, font_size=11)

    doc.add_paragraph()

    # Testate
    p_test = doc.add_paragraph()
    r = p_test.add_run("Testate presenti: ")
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r.bold = True
    r2 = p_test.add_run(", ".join(testate))
    r2.font.name = "Arial"
    r2.font.size = Pt(10)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # SEZIONI ANALITICHE
    # ════════════════════════════════════════════════════════════════
    sezioni = [
        ("1. Profilo Mediatico (Media Narrative Profile)", "sezione1_profilo"),
        ("2. Interviste e Presenza dei Vertici",           "sezione2_vertici"),
        ("3. Temi Longevi (Long Running Issues)",          "sezione3_temi_longevi"),
        ("4. Notizie Finanziarie e Corporate",             "sezione4_finanziario"),
        ("5. Cambi di Management e Governance",            "sezione5_governance"),
        ("6. Focus Territoriale e Temi Locali",            "sezione6_territoriale"),
        ("7. Criticita Reputazionali",                     "sezione7_criticita"),
        ("8. Analisi del Sentiment",                       "sezione8_sentiment"),
        ("9. Comunicazione Istituzionale",                 "sezione9_istituzionale"),
        ("10. Sintesi Strategica Finale",                  "sezione10_sintesi"),
    ]

    for titolo_sezione, key in sezioni:
        set_heading_style(1, titolo_sezione)
        add_separator()
        testo = sections.get(key, "Dati non disponibili per questa sezione.")
        add_body(testo)
        doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════
    # APPENDICE — LISTA ARTICOLI
    # ════════════════════════════════════════════════════════════════
    doc.add_page_break()
    set_heading_style(1, "Appendice — Elenco Articoli")

    tbl2 = doc.add_table(rows=1, cols=5)
    tbl2.style = "Table Grid"
    for i, h in enumerate(["Testata", "Data", "Autore", "Titolo", "AVE (€)"]):
        shade_cell(tbl2.rows[0].cells[i], "1A2E4A")
        cell_text(tbl2.rows[0].cells[i], h, bold=True, color=C_WHITE, font_size=8)

    for idx, a in enumerate(articles):
        row = tbl2.add_row()
        fill = "F5F5F5" if idx % 2 == 0 else "FFFFFF"
        vals = [
            a.get("testata",""),
            str(a.get("data",""))[:10],
            (a.get("giornalista") or a.get("autore",""))[:30],
            (a.get("titolo",""))[:60],
            str(a.get("ave","")) or "N/D",
        ]
        for i, v in enumerate(vals):
            shade_cell(row.cells[i], fill)
            cell_text(row.cells[i], v, font_size=8)

    # ── Nota finale ──
    doc.add_paragraph()
    p_nota = doc.add_paragraph()
    p_nota.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_nota.add_run(f"Fine documento — MAIM Intelligence & Media Monitoring | {periodo} | {n_art} articoli elaborati")
    r.font.name   = "Arial"
    r.font.size   = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ── Salva ──
    filename = f"report_{date_from}_{uuid.uuid4().hex[:8]}.docx"
    filepath = REPORTS_DIR / filename
    doc.save(str(filepath))
    return filepath


def generate_word_report(question: str, keywords: list[str], client_data: dict | None,
                          date_from: str, date_to: str, db_filters: dict) -> str:
    """
    Orchestratore del report Word:
    1. Fetch tutti gli articoli pertinenti (metadati)
    2. GPT-4o genera le 10 sezioni in JSON
    3. python-docx costruisce il Word
    4. Restituisce link di download
    """
    # Fetch articoli — usiamo FIELDS_REPORT (metadati + occhiello, no testo integrale)
    # Nessun limite: prendiamo tutti gli articoli disponibili
    articles = _run_query(FIELDS_REPORT, date_from, date_to, keywords, db_filters, limit=500)

    if not articles:
        periodo = f"del {date_from}" if date_from == date_to else f"dal {date_from} al {date_to}"
        return f"Nessun articolo trovato {periodo} per generare il report."

    n = len(articles)
    topic_str = ", ".join(keywords) if keywords else "generale"
    periodo   = f"del {date_from}" if date_from == date_to else f"dal {date_from} al {date_to}"

    # Avviso all'utente che stiamo lavorando
    print(f"[REPORT] {n} articoli trovati — chiamo GPT-4o per analisi...")

    try:
        sections = _call_gpt_for_report(articles, keywords, client_data, date_from, date_to)
    except json.JSONDecodeError as e:
        return f"Errore nel parsing della risposta GPT-4o: {e}. Riprova."
    except openai.RateLimitError:
        return "Limite token OpenAI raggiunto. Aspetta 60 secondi e riprova."
    except Exception as e:
        return f"Errore nella generazione del report: {e}"

    try:
        filepath = _build_word_document(sections, articles, keywords, client_data, date_from, date_to)
    except Exception as e:
        return f"Errore nella costruzione del documento Word: {e}"

    download_url = f"{BASE_URL}/static/reports/{filepath.name}"

    return (
        f"**Report generato con successo.**\n\n"
        f"- Periodo: {periodo}\n"
        f"- Topic: {topic_str}\n"
        f"- Articoli analizzati: {n}\n"
        f"- Testate: {len(set(a.get('testata','') for a in articles))}\n\n"
        f"**[Scarica il report Word]({download_url})**"
    )


# ═══════════════════════════════════════════════════════════════════════════════
# SYSTEM PROMPT PER OPENAI (chat normale)
# ═══════════════════════════════════════════════════════════════════════════════
_INTENT_GUIDE = {
    "leggi":   "MODALITA' LETTURA: restituisci il TESTO INTEGRALE dell'articolo piu' pertinente.\nIntestazione: **[TESTATA] — [TITOLO]** | di [AUTORE] | [Leggi](LINK)",
    "rischio": "MODALITA' RISK ALERT: elenca articoli con rischio reputazionale o politico elevato.\nPer ognuno: Titolo | Testata | Autore | Tipo rischio | Motivazione | Link.",
    "analisi": "MODALITA' ANALISI: estrai temi, posizioni, tendenze. Confronta fonti. Sintesi finale 3-5 righe.",
    "report":  "MODALITA' REPORT TESTO: rassegna professionale con sezioni chiare. Titoli principali con testata, autore e link, poi approfondimento per tema.",
    "generico":"Rispondi con i dati disponibili. Cita sempre testata, autore e link [Leggi](URL) per ogni articolo.",
}

def build_system_prompt(intent: str, client_data: dict | None,
                         date_from: str, date_to: str) -> str:
    periodo = f"del {date_from}" if date_from == date_to else f"dal {date_from} al {date_to}"
    client_block = ""
    if client_data:
        client_block = (
            f"\nCLIENTE: **{client_data.get('name')}** | "
            f"Keywords: {client_data.get('keywords','—')} | "
            f"Settore: {client_data.get('semantic_topic','—')}\n"
        )
    return (
        "Sei SPIZ, analista professionale di rassegna stampa.\n"
        "Per ogni articolo cita SEMPRE: Testata, Autore, Titolo, Link [Leggi](URL).\n"
        "Non inventare dati. Sii preciso e diretto.\n\n"
        f"PERIODO: {periodo}{client_block}\n"
        f"{_INTENT_GUIDE.get(intent, _INTENT_GUIDE['generico'])}"
    )

# ═══════════════════════════════════════════════════════════════════════════════
# FUNZIONE PRINCIPALE
# ═══════════════════════════════════════════════════════════════════════════════
def ask_spiz(
    question:   str,
    session_id: str = "default",
    client_id:  str | None = None,
) -> str:

    q_low = question.strip().lower()

    # ── Comandi speciali ──────────────────────────────────────────────────────
    if q_low in ("reset", "/reset", "nuova conversazione"):
        reset_session(session_id)
        return "Conversazione azzerata."

    if q_low in ("clienti", "/clienti", "lista clienti"):
        clients = list_clients()
        if not clients:
            return "Nessun cliente nel database."
        rows_txt = "\n".join(
            f"- **{c['name']}** (ID: `{c['id']}`)\n"
            f"  Keywords: {c.get('keywords','—')} | Settore: {c.get('semantic_topic','—')}"
            for c in clients
        )
        return f"**Clienti registrati ({len(clients)}):**\n\n{rows_txt}"

    if any(t in q_low for t in ("in memoria", "hai in memoria", "nel contesto", "quanti ne hai")):
        n_turni = len(get_history(session_id)) // 2
        if n_turni == 0:
            return "Non ho ancora nessuno scambio in memoria. La memoria si costruisce man mano che fai domande."
        return (
            f"In questa sessione ho risposto a **{n_turni} domande**.\n\n"
            "Ogni domanda interroga il database in tempo reale."
        )

    # ── Carica cliente ────────────────────────────────────────────────────────
    client_data = None
    if client_id:
        client_data = load_client(client_id)
        if not client_data:
            return f"Cliente '{client_id}' non trovato nel database."

    # ── Date, intent, filtri ──────────────────────────────────────────────────
    date_from, date_to = date_range_from_question(question)
    db_filters         = extract_db_filters(question)
    intent             = classify_intent(question)

    # ── Keyword topic ─────────────────────────────────────────────────────────
    if client_data:
        keywords = parse_keywords(client_data.get("keywords", ""))
        topic    = client_data.get("semantic_topic", "")
        if topic:
            keywords.append(topic.lower())
    else:
        keywords = topic_keywords_from_question(question)

    # ══════════════════════════════════════════════════════════════════════════
    # PERCORSO A — Report Word (fetch massivo + GPT-4o + python-docx)
    # ══════════════════════════════════════════════════════════════════════════
    if intent == "report_word":
        return generate_word_report(question, keywords, client_data,
                                     date_from, date_to, db_filters)

    # ══════════════════════════════════════════════════════════════════════════
    # PERCORSO B — DB diretto, 0 token OpenAI
    # ══════════════════════════════════════════════════════════════════════════
    if intent in ("conta", "autore", "fonte", "ave"):
        rows = _run_query(FIELDS_LIGHT, date_from, date_to, keywords, db_filters, limit=1000)
        if not rows:
            periodo = f"del {date_from}" if date_from == date_to else f"dal {date_from} al {date_to}"
            return f"Nessun articolo trovato {periodo}."
        if intent == "conta":  return resolve_conta(rows, question, date_from, date_to)
        if intent == "autore": return resolve_autore(rows, date_from, date_to)
        if intent == "fonte":  return resolve_fonte(rows, date_from, date_to)
        if intent == "ave":    return resolve_ave(rows, date_from, date_to)

    # ══════════════════════════════════════════════════════════════════════════
    # PERCORSO C — OpenAI (analisi, report testuale, lettura, generico)
    # ══════════════════════════════════════════════════════════════════════════
    articles = fetch_full_articles(keywords, date_from, date_to,
                                    extra_filters=db_filters, limit=8)
    if not articles:
        periodo = f"del {date_from}" if date_from == date_to else f"dal {date_from} al {date_to}"
        chi     = f" per il cliente '{client_data['name']}'" if client_data else ""
        return f"Nessun articolo trovato {periodo}{chi}."

    history    = trim_history(get_history(session_id), budget=2_000)
    ctx_budget = MAX_CONTEXT_TOKENS - MAX_RESPONSE_TOKENS - 1_500
    db_context, _ = build_context(articles, budget=ctx_budget)

    model         = MODEL_SMART if intent in ("analisi", "report") else MODEL_FAST
    system_prompt = build_system_prompt(intent, client_data, date_from, date_to)
    user_message  = f"DATABASE:\n{db_context}\n\nRICHIESTA: {question}"

    messages = [
        {"role": "system", "content": system_prompt},
        *history,
        {"role": "user",   "content": user_message},
    ]

    try:
        response = openai.chat.completions.create(
            model=model, messages=messages, temperature=0,
            max_tokens=MAX_RESPONSE_TOKENS,
        )
        answer = response.choices[0].message.content
        save_turn(session_id, question, answer)
        return answer
    except openai.RateLimitError:
        return "Limite token OpenAI raggiunto. Aspetta 60 secondi e riprova."
    except openai.APIError as e:
        return f"Errore API OpenAI: {e}"
    except Exception as e:
        return f"Errore inatteso: {e}"


# ─── TEST DA TERMINALE ──────────────────────────────────────────────────────────
if __name__ == "__main__":
    import sys
    client_arg = sys.argv[1] if len(sys.argv) > 1 else None
    session    = "test"
    if client_arg:
        print(f"Modalita' cliente: {client_arg}\n")
    print("SPIZ pronto. Comandi: 'reset', 'clienti'. Ctrl+C per uscire.\n")
    while True:
        try:
            q = input("Tu: ").strip()
            if not q:
                continue
            print(f"\nSPIZ: {ask_spiz(q, session_id=session, client_id=client_arg)}\n")
        except KeyboardInterrupt:
            print("\nArrivederci.")
            break